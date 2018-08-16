from docx import Document

document = Document("doc.docx")


print(document.tables)

raise SystemExit("My Temp Break Point")

for table in document.tables:
    print("############################\n")
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                print(para.text)
    print("##############################\n")
